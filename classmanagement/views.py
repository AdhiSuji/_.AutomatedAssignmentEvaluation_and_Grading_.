#View Handling & Auth
from django.shortcuts import render, redirect, get_object_or_404  
from django.contrib.auth.decorators import login_required, user_passes_test  
from django.contrib.auth import authenticate, login, logout, get_user_model  
from django.contrib import messages  

from collections import defaultdict

#Email
from django.core.mail import send_mail  
#Logging & JSON
import json
#Database & Time  
from django.utils.timezone import now  
from django.utils import timezone  
#HTTP & Queries 
from django.http import JsonResponse  
from django.db.models import Q  
#Pagination
from django.core.paginator import Paginator  
#MOdels & Forms
from .models import ( CustomUser, Classroom, Assignment, Submission, Enrollment, Performance, StudentProfile, TeacherProfile, PrivateMessage, QueryMessage , Notification )  
from .forms import ( AssignmentForm, SubmissionForm, ClassForm, StudentProfileForm, TeacherProfileForm  )  


def is_admin(user):
    return user.is_authenticated and user.role == 'admin'

def is_teacher(user):
    return user.is_authenticated and user.role == 'teacher'

def is_student(user):
    return user.is_authenticated and user.role == 'student'

def profile_view(request):
    profile = StudentProfile.objects.get(user=request.user)
    print(profile.profile_pic)
    print(profile.profile_pic.url)
    
    return render(request, 'profile.html', {'profile': profile})

 #VIEWS             

def home_view(request):
    return render(request, 'home.html')

def privacy_policy(request):
    return render(request, 'privacy_policy.html')


def contact_us(request):
    if request.method == "POST":
        name = request.POST.get('name')
        email = request.POST.get('email')
        message = request.POST.get('message')

        # Sending email (you need to configure email settings in Django)
        send_mail(
            f"Contact Form Submission from {name}",
            message,
            email,
            ['submitech1@gmail.com'],  # Replace with your email
            fail_silently=False,
        )

        messages.success(request, "Your message has been sent!")
    
    return render(request, 'contact_us.html')

# REGISTRATION      
def register(request):
    if request.method == "POST":
        first_name = request.POST.get("first_name")
        last_name = request.POST.get("last_name")
        email = request.POST.get("email")
        role = request.POST.get("role")
        password = request.POST.get("password")
        confirm_password = request.POST.get("confirm_password")

        if password != confirm_password:
            messages.error(request, "Passwords do not match!")
            return render(request, "register.html", {"form_data": request.POST})

        if CustomUser.objects.filter(email=email).exists():
            messages.error(request, "Email is already registered!")
            return render(request, "register.html", {"form_data": request.POST})

        user = CustomUser.objects.create_user(
            username=email,
            email=email,
            first_name=first_name,
            last_name=last_name,
            password=password,
            role=role
        )

        login(request, user)

        if role == "student":
            # Create StudentProfile + Performance
            student_profile = StudentProfile.objects.create(student=user)
            Performance.objects.create(student=student_profile)

            messages.success(request, "Student registered successfully!")
            return redirect("student_profile")

        elif role == "teacher":
            # Create TeacherProfile
            teacher_profile = TeacherProfile.objects.create(teacher=user)
            messages.success(request, "Teacher registered successfully!")
            return redirect('teacher_create_class')

        messages.error(request, "Invalid role! Contact admin.")
        return redirect('register')

    return render(request, "register.html")


#LOGIN           
def user_login(request):
    if request.method == "POST":
        email = request.POST.get("email")
        password = request.POST.get("password")

        user = authenticate(request, email=email, password=password)

        if not user:
            messages.error(request, "Invalid email or password!")
            return redirect("login")

        login(request, user)
        messages.success(request, "Login successful!")

        # STUDENT LOGIN
        if user.role == "student":
            # Get or Create StudentProfile
            student_profile, created = StudentProfile.objects.get_or_create(student=user)

            # Create performance record if new profile was created
            if created:
                Performance.objects.create(student=student_profile)

            # ✅ FIX: Access classrooms instead of assigned_class
            joined_classes = student_profile.joined_classes.all()

            if joined_classes.exists():
                first_class = joined_classes.first()  # Grab the first class
                return redirect("student_dashboard", class_id=first_class.id)
            else:
                messages.warning(request, "No class assigned! Contact your teacher.")
                return redirect("student_profile")

        # TEACHER LOGIN
        elif user.role == "teacher":
            # Get or Create TeacherProfile (avoid duplicates)
            teacher_profile, created = TeacherProfile.objects.get_or_create(teacher=user)

            # ✅ Corrected: Use teacher_profile instead of user
            teacher_classes = Classroom.objects.filter(teacher=teacher_profile)

            if teacher_classes.exists():
                class_obj = teacher_classes.first()
                return redirect('teacher_dashboard', class_id=class_obj.id)
            else:
                return redirect('teacher_profile')

        # ADMIN LOGIN
        elif user.role == "admin":
            return redirect("admin_dashboard")

        # Unknown role
        messages.error(request, "Invalid role! Contact admin.")
        return redirect("login")

    return render(request, "login.html")

def user_logout(request):
    logout(request)
    return redirect('login')

def student_default_profile():
    return 'default_folder/default_student.jpg'  # path inside media folder

def teacher_default_profile():
    return 'default_folder/default_teacher.jpg'  # path inside media folder

@login_required
@user_passes_test(is_teacher)
def teacher_profile(request):
    teacher_profile = get_object_or_404(TeacherProfile, teacher=request.user)

    # Fetch all classrooms assigned to this teacher
    created_classes = Classroom.objects.filter(teacher=teacher_profile)

    # Initialize search query
    query = request.GET.get('q')

    # Fetch all students assigned to this teacher's classes
    students = StudentProfile.objects.filter(joined_classes__in=created_classes).distinct()

    if query:
        students = students.filter(
            Q(name__icontains=query) |
            Q(student__email__icontains=query)
        )

    # Paginate the students list
    paginator = Paginator(students, 10)  # Show 10 students per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    # If no classrooms are created, show a message
    if not created_classes.exists():
        messages.info(request, "You haven't created any classrooms yet.")

    # Prepare the form for editing teacher profile (if needed)
    form = TeacherProfileForm(instance=teacher_profile)

    return render(request, 'teacher_profile.html', {
        'teacher_profile': teacher_profile,
        'created_classes': created_classes,
        'page_obj': page_obj,
        'query': query,
        'form': form,  # Ensure form is passed if editing is required
    })


@login_required
@user_passes_test(is_student)
def student_profile(request):
    student_user = request.user
    student_profile, created = StudentProfile.objects.get_or_create(student=student_user)

    if request.method == 'POST':
        form = StudentProfileForm(request.POST, request.FILES, instance=student_profile)
        if form.is_valid():
            form.save()
            messages.success(request, "Profile updated successfully!")
            return redirect('student_profile')
    else:
        form = StudentProfileForm(instance=student_profile)

    joined_classes = student_profile.joined_classes.all()
    assignments = Assignment.objects.filter(joined_classes__in=joined_classes)
    teachers = TeacherProfile.objects.filter(classrooms__in=joined_classes).distinct()
    performance, _ = Performance.objects.get_or_create(student=student_profile)

    classroom = joined_classes.first() if joined_classes.exists() else None
    teacher = teachers.first()  # Get the first teacher if multiple exist

    teacher_id = teacher.teacher.id if teacher else None  # ✅ Get correct teacher ID
    student_id = student_profile.student.id if student_profile else None  # ✅ Get correct student ID

    context = {
        'student_profile': student_profile,
        'joined_classes': joined_classes,
        'assignments': assignments,
        'teachers': teachers,
        'performance': performance,
        'form': form,
        'classroom': classroom,
        'teacher_id': teacher_id,
        'student_id': student_id,
    }

    return render(request, 'student_profile.html', context)


@login_required
def teacher_dashboard(request, class_id):
    """Teacher's Dashboard: Displays student performance for a selected class."""
    
    teacher = get_object_or_404(TeacherProfile, teacher=request.user)

    # ✅ Get all classes of this teacher
    classes = Classroom.objects.filter(teacher=teacher)

    # ✅ Fetch the currently selected class
    current_class = get_object_or_404(Classroom, id=class_id, teacher=teacher)

    # ✅ Fetch students in this class
    students = StudentProfile.objects.filter(joined_classes=current_class)

    # ✅ Fetch assignments only from this class
    assignments = Assignment.objects.filter(joined_classes=current_class)

    # ✅ Fetch submissions only from this class
    submissions = Submission.objects.filter(assignment__joined_classes=current_class)

    performance_data = []
    student_performance_list = []
    class_avg_performance = {}

    for student in students:
        student_performance = {
            "student": student,
            "assignments": [],
            "total_score": 0
        }

        # ✅ Fetch submissions only for this student in this class
        student_submissions = submissions.filter(student=student)
        for submission in student_submissions:
            student_performance["assignments"].append({
                "assignment": submission.assignment,
                "is_late": submission.is_late,
                "plagiarism_score": submission.plagiarism_score,
                "text_similarity_score": submission.text_similarity_score,
                "image_similarity_score": submission.image_similarity_score,
                "provisional_marks": float(submission.provisional_marks), 
                "final_marks": float(submission.final_marks),  # ✅ Add this
                "grade_before_plagiarism": submission.grade_before_plagiarism,  # ✅ Add this
                "grade_after_plagiarism": submission.grade_after_plagiarism,  # ✅ Add this
                "feedback_before_plagiarism": submission.feedback_before_plagiarism,  # ✅ Add this
                "feedback_after_plagiarism": submission.feedback_after_plagiarism,  # ✅ Add this

            })

            student_performance_list.append({
                "student": student.student.username,  # Use full name if needed
                "assignment": submission.assignment.title,
                "marks": float(submission.final_marks)  # Convert Decimal to float here too
            })

            student_performance["total_score"] += float(submission.final_marks)  # Sum as float

            # ✅ Compute class average per assignment
            if submission.assignment.title not in class_avg_performance:
                class_avg_performance[submission.assignment.title] = []
            class_avg_performance[submission.assignment.title].append(float(submission.final_marks))  # Convert to float

        performance_data.append(student_performance)

    # ✅ Compute class averages
    for assignment in class_avg_performance:
        scores = class_avg_performance[assignment]
        class_avg_performance[assignment] = sum(scores) / len(scores) if scores else 0

    # ✅ Fetch top 3 students based on performance
    sorted_performance_data = sorted(performance_data, key=lambda x: x["total_score"], reverse=True)
    top_students = sorted_performance_data[:3]

    # ✅ Convert data to JSON for JavaScript charts
    student_performance_json = json.dumps(student_performance_list)
    avg_class_performance_json = json.dumps(class_avg_performance)

    return render(request, 'teacher_dashboard.html', {
        'classes': classes,
        'current_class': current_class,
        'students': students,
        'assignments': assignments,
        'performance_data': performance_data,
        'top_students': top_students,
        'student_performance_json': student_performance_json,
        'avg_class_performance_json': avg_class_performance_json
    })



@login_required
def student_dashboard(request, class_id=None):
    student_profile = get_object_or_404(StudentProfile, student=request.user)
    classroom = get_object_or_404(Classroom, id=class_id)

    # Which marks to display: 'provisional' or 'final'
    marks_type = request.GET.get('marks_type', 'provisional')  # default to 'provisional'

    # Mapping marks_type to the correct grade field
    grade_field_mapping = {
        'provisional': 'grade_before_plagiarism',
        'final': 'grade_after_plagiarism',
    }
    grade_field = grade_field_mapping.get(marks_type, 'grade_after_plagiarism')

    # Fetch all submissions by the student for this class
    submissions = Submission.objects.filter(
        student=student_profile,
        assignment__joined_classes=classroom
    ).order_by('-submitted_at')

    # Preparing student performance data for chart/analysis
    student_performance = [{
        "assignment": s.assignment.title,
        "marks": float(getattr(s, f'{marks_type}_marks')),  # e.g., provisional_marks or final_marks
        "grade": getattr(s, grade_field),             # e.g., grade_before_plagiarism or grade_after_plagiarism
        "feedback": s.feedback_after_plagiarism or s.feedback_before_plagiarism
    } for s in submissions]

    # Fetch all students in the classroom
    student_enrollments = Enrollment.objects.filter(classroom=classroom, role='student')
    all_students = [enrollment.student for enrollment in student_enrollments]

    # Calculate overall class performance
    class_performance = defaultdict(list)
    student_scores = {}

    submissions_for_class = Submission.objects.filter(
        assignment__joined_classes=classroom
    ).select_related('student', 'assignment')

    for sub in submissions_for_class:
        marks = getattr(sub, f'{marks_type}_marks')
        class_performance[sub.assignment.title].append(marks)

        student_name = f"{sub.student.student.first_name} {sub.student.student.last_name}".strip()
        if student_name in student_scores:
            student_scores[student_name]['overall_score'] += marks
        else:
            student_scores[student_name] = {
                "student_name": student_name,
                "overall_score": marks
            }

    # Find top 3 performers based on overall marks
    top_performers = sorted(student_scores.values(), key=lambda x: x['overall_score'], reverse=True)[:3]

    # Render the dashboard
    return render(request, 'student_dashboard.html', {
        'submissions': submissions,
        'student_performance_json': json.dumps(student_performance),
        'top_performers': top_performers,
        'current_class': classroom,
        'student': student_profile,
        'marks_type': marks_type,
    })


#ADMIN VIEWS (Optional)   
@login_required
@user_passes_test(is_admin)
def admin_dashboard(request):
    teachers = CustomUser.objects.filter(role='teacher')
    students = CustomUser.objects.filter(role='student')
    classes = Classroom.objects.all()

    return render(request, 'admin_dashboard.html', {
        'teachers': teachers,
        'students': students,
        'classes': classes
    })


@login_required
@user_passes_test(is_teacher)
def create_class(request):
    if request.method == "POST":
        form = ClassForm(request.POST)
        if form.is_valid():
            new_class = form.save(commit=False)
            teacher_profile = TeacherProfile.objects.get(teacher=request.user)
            new_class.teacher = teacher_profile
            new_class.save()
            messages.success(request, "Class created successfully!")
            return redirect('teacher_dashboard',class_id=new_class.id)

    else:
        form = ClassForm()

    return render(request, 'create_class.html', {'form': form})


@login_required
@user_passes_test(is_teacher)
def teacher_classes(request):
    teacher_profile = get_object_or_404(TeacherProfile, teacher=request.user)
    classes = Classroom.objects.filter(teacher=teacher_profile)
    context = {
        'classes': classes
    }

    return render(request, 'teacher_classes.html', context)


@login_required
@user_passes_test(is_student)
def enroll_students(request, class_id):
    print("DEBUG: Is user authenticated?", request.user.is_authenticated)
    print("DEBUG: Logged-in user:", request.user)
    print("DEBUG: User groups:", request.user.groups.all())

    if request.method == "POST":
        reference_id = request.POST.get("reference_id")
        class_id = request.POST.get("selected_class")

        teacher = TeacherProfile.objects.filter(reference_id=reference_id).first()
        if not teacher:
            messages.error(request, "Invalid Teacher Reference ID.")
            return redirect('enroll_student')

        selected_class = Classroom.objects.filter(id=class_id, teacher=teacher).first()
        if not selected_class:
            messages.error(request, "Invalid Class Selection.")
            return redirect('enroll_student')

        student_profile = get_object_or_404(StudentProfile, student=request.user)

        if selected_class in student_profile.joined_classes.all():
            messages.warning(request, "You are already enrolled in this class.")
        else:
            student_profile.joined_classes.add(selected_class)
            messages.success(request, f"Successfully enrolled in {selected_class.name}.")

        return redirect('student_dashboard')

    return render(request, 'enroll_student.html')

@login_required
@user_passes_test(is_teacher)
def view_enrolled_students(request, class_id):
    classroom = get_object_or_404(Classroom, id=class_id, teacher__teacher=request.user)
    students = classroom.joined_students.all()  # ✅ Correct way to get enrolled students

    return render(request, 'view_students.html', {'classroom': classroom, 'students': students})


def teacher_list(request):
    teachers = CustomUser.objects.filter(role="teacher")
    return render(request, "teacher_list.html", {"teachers": teachers})

@login_required
def add_teacher(request):
    if request.method == 'POST':
        teacher_reference_id = request.POST.get('teacher_reference_id')
        joined_classes_id = request.POST.get('joined_classes')

        if not teacher_reference_id or not joined_classes_id:
            messages.error(request, "Please enter both Teacher Reference ID and select a class.")
            return redirect('student_profile')

        try:
            # Get class by assigned_class_id
            classroom = Classroom.objects.get(id=joined_classes_id, teacher__reference_id=teacher_reference_id)

            # Get student profile
            student_profile = request.user.student_profile

            # Add the class to the student’s profile (ManyToManyField)
            student_profile.joined_classes.add(classroom)

            messages.success(request, f"You have successfully joined {classroom.name}!")
            return redirect('student_profile')

        except Classroom.DoesNotExist:
            messages.error(request, "Classroom not found for the given reference ID.")
            return redirect('student_profile')

    else:
        return redirect('student_profile')

from django.shortcuts import get_object_or_404
from django.http import JsonResponse

def get_teacher_classes(request):
    reference_id = request.GET.get('reference_id')
    print("DEBUG: Got reference_id:", reference_id)

    if not reference_id:
        return JsonResponse({'error': 'No reference ID provided'}, status=400)

    teacher = get_object_or_404(TeacherProfile, reference_id=reference_id)
    classes = Classroom.objects.filter(teacher=teacher)
    print("DEBUG: Found classes:", list(classes))

    class_list = [
        {'id': cls.id, 'name': cls.name}
        for cls in classes
    ]

    return JsonResponse({'classes': class_list})


def join_class(request):
    if request.method == 'POST':
        teacher_reference_id = request.POST.get('teacher_reference_id')
        joined_classes_id = request.POST.get('joined_classes')

        print("Received POST Data:", request.POST)  # Debugging output

        if not teacher_reference_id or not joined_classes_id:
            messages.error(request, "Please enter both Teacher Reference ID and select a class.")
            return redirect('join_class')  # Redirect back to the form
        
        try:
            teacher = TeacherProfile.objects.get(reference_id=teacher_reference_id)
            classroom = Classroom.objects.get(id=joined_classes_id)

            student_profile, created = StudentProfile.objects.get_or_create(student=request.user)
            if classroom in student_profile.joined_classes.all():
                messages.warning(request, 'You have already joined this class.')
                return redirect('student_profile')

            student_profile.joined_classes.add(classroom)
            student_profile.save()

            messages.success(request, 'Successfully joined the class!')
            return redirect('student_profile')

        except TeacherProfile.DoesNotExist:
            messages.error(request, 'Teacher not found.')
        except Classroom.DoesNotExist:
            messages.error(request, 'Class not found.')
        except Exception as e:
            messages.error(request, f"Unexpected error: {e}")
            print(f"Unexpected error: {e}")  # Print error for debugging

    return redirect('join_class')


from django.contrib.auth.decorators import login_required
from django.shortcuts import get_object_or_404, redirect
from django.contrib import messages
from .models import Classroom, Enrollment

@login_required
def delete_class(request, class_id):
    # Get the classroom object
    classroom = get_object_or_404(Classroom, id=class_id)

    # Check if the user is the teacher of the class
    if request.user == classroom.teacher.teacher:  # Access the teacher's user
        if request.method == "POST":
            # Delete all enrollments for this class first
            Enrollment.objects.filter(classroom=classroom).delete()
            # Now delete the classroom
            classroom.delete()
            messages.success(request, f"Class '{classroom.name}' deleted successfully!")
            return redirect('teacher_profile')

    # Check if the user is a student enrolled in the class
    elif Enrollment.objects.filter(student=request.user, classroom_id=class_id).exists():
        if request.method == "POST":
            # Remove the student from the class
            enrollment = Enrollment.objects.get(student=request.user, classroom_id=class_id)
            enrollment.delete()
            messages.success(request, f"You have left the class '{classroom.name}'.")
            return redirect('student_profile')

    # If the user is neither the teacher nor a student in the class, show an error message
    messages.error(request, "You do not have permission to perform this action.")
    return redirect('student_profile')

@login_required
@user_passes_test(is_teacher)
def give_assignment(request, class_id):
    teacher_profile = get_object_or_404(TeacherProfile, teacher=request.user)

    classroom = get_object_or_404(Classroom, id=class_id)

    if request.method == "POST":
        form = AssignmentForm(request.POST,request.FILES)
        if form.is_valid():
            assignment = form.save(commit=False)
            assignment.teacher = teacher_profile  # ✅ Now it's the correct object type
            assignment.joined_classes = classroom  # ✅ Assign class based on class_id
            assignment.save()
            messages.success(request, "Assignment given successfully!")
            return redirect('teacher_dashboard', class_id=class_id)
    else:
        form = AssignmentForm()

    # Optional: fetch existing assignments for the class
    given_assignments = Assignment.objects.filter(teacher=teacher_profile)

    paginator = Paginator(given_assignments, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    return render(request, 'give_assignment.html', {
        'form': form,
        'page_obj': page_obj
    })


@login_required
def given_assignment(request, class_id):
    class_obj = get_object_or_404(Classroom, id=class_id)
    assignments = Assignment.objects.filter(joined_classes=class_obj)

    return render(request, "given_assignment.html", {
        "assignments": assignments,
        "class_obj": class_obj
    })



@login_required
@user_passes_test(is_teacher)
def view_submissions(request, classroom_id=None, student_id=None):
    teacher_profile = get_object_or_404(TeacherProfile, teacher=request.user)
    submissions = Submission.objects.filter(assignment__teacher=teacher_profile)

    classroom = None  # Initialize classroom as None

    if classroom_id:
        classroom = get_object_or_404(Classroom, id=classroom_id)  # ✅ Fetch classroom object
        submissions = submissions.filter(assignment__joined_classes__id=classroom_id)

    if student_id:
        submissions = submissions.filter(student_id=student_id)

    context = {
        'submissions': submissions,
        'classroom': classroom,  # ✅ Pass classroom to template
        'classroom_id': classroom_id,
        'student_id': student_id,
    }

    return render(request, 'view_submissions.html', context)



# Create assignment view
def create_assignment(request):
    if request.method == 'POST':
        form = AssignmentForm(request.POST, request.FILES)
        if form.is_valid():
            assignment = form.save()

            file_path = assignment.model_answer_file.path
            print(f"Assignment file saved at: {file_path}")
            text = ""

            # ✅ TEXT Extraction
            try:
                text = extract_text_from_pdf(file_path)  # Calls the function to extract text
                assignment.extracted_text = text  # Save the extracted text
                assignment.save(update_fields=['extracted_text'])
            except Exception as e:
                print("Text Extraction Failed:", e)

            print("Extracted Text:\n", text)

            # ✅ IMAGE Extraction (PDF only)
            try:
                if file_path.endswith('.pdf'):
                    image_paths = extract_images_from_pdf(file_path, assignment.id)  # Extract images from PDF
                    # Optionally save the first image as the preview
                    if image_paths:
                        assignment.extracted_image.name = image_paths[0]  # Save first image as the extracted image
                        assignment.save(update_fields=['extracted_image'])
            except Exception as e:
                print("Diagram Extraction Failed:", e)

            messages.success(request, "Assignment uploaded and model answer processed.")
            return redirect('teacher_dashboard')
    else:
        form = AssignmentForm()

    return render(request, 'give_assignment.html', {'form': form})


@login_required
@user_passes_test(is_student)
def submit_assignment(request, assignment_id):
    assignment = get_object_or_404(Assignment, id=assignment_id)
    student_profile = request.user.student_profile
    student_classes = student_profile.joined_classes.all()

    # Ensure that the student is in the class related to the assignment
    if not student_classes.filter(id=assignment.joined_classes.id).exists():
        messages.error(request, "You are not enrolled in this class.")
        return redirect('student_dashboard')

    # Fetch assignments only for the student's classes, filter by the current assignment class
    query = request.GET.get('q')
    assignments = Assignment.objects.filter(joined_classes__in=student_classes, id=assignment.id).distinct()
    
    if query:
        assignments = assignments.filter(title__icontains=query)

    paginator = Paginator(assignments, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    if request.method == 'POST':
        form = SubmissionForm(request.POST, request.FILES)
        if form.is_valid():
            submission = form.save(commit=False)
            submission.student = student_profile
            submission.assignment = assignment
            submission.submitted_at = timezone.now()
            submission.is_late = timezone.now() > assignment.due_date
            submission.save()

            evaluate_submission(submission.id)

            messages.success(request, 'Assignment submitted and evaluated successfully!')
            return redirect('student_dashboard')
    else:
        form = SubmissionForm()

    return render(request, 'submit_assignment.html', {
        'form': form,
        'assignment': assignment,
        'assignments': assignments,
        'page_obj': page_obj,
        'query': query
    })


import os
import cv2
import fitz  # PyMuPDF
import logging
import numpy as np
from PIL import Image
from textblob import TextBlob

from django.conf import settings
from django.utils import timezone
from django.core.files import File
from django.db import transaction

import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer

from sentence_transformers import SentenceTransformer
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

import pytesseract
# ------------------- Setup --------------------
logger = logging.getLogger(__name__)
pytesseract.pytesseract.tesseract_cmd = r'C:\Tesseract-OCR\tesseract.exe'

nltk.download('punkt')
nltk.download('stopwords')
nltk.download('wordnet')

TEXT_MODEL = SentenceTransformer('all-MiniLM-L6-v2')
IMG_MODEL = SentenceTransformer('paraphrase-MiniLM-L6-v2')

STUDENT_IMG_DIR = os.path.join(settings.MEDIA_ROOT, 'extracted_student_images')
TEACHER_IMG_DIR = os.path.join(settings.MEDIA_ROOT, 'extracted_teacher_images')
os.makedirs(STUDENT_IMG_DIR, exist_ok=True)
os.makedirs(TEACHER_IMG_DIR, exist_ok=True)

import os
import fitz
import cv2
import pytesseract
import numpy as np
import docx
from PIL import Image
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.utils import timezone
from django.contrib.auth.decorators import login_required, user_passes_test
from sklearn.metrics.pairwise import cosine_similarity
from sentence_transformers import SentenceTransformer
from .models import Assignment, Submission
from .forms import SubmissionForm
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from textblob import TextBlob
from sklearn.feature_extraction.text import TfidfVectorizer

# Global Text Model
TEXT_MODEL = SentenceTransformer('paraphrase-MiniLM-L6-v2')

# ---------------------- Utility Functions ----------------------

def is_pdf(path):
    return path.lower().endswith('.pdf')

def is_docx(path):
    return path.lower().endswith('.docx')

def extract_text_from_pdf(path):
    doc = fitz.open(path)
    return "".join(page.get_text() for page in doc)

def extract_text_from_docx(path):
    doc = docx.Document(path)
    return "\n".join(para.text for para in doc.paragraphs)

def extract_text(path):
    if is_pdf(path):
        return extract_text_from_pdf(path)
    elif is_docx(path):
        return extract_text_from_docx(path)
    return ""

def preprocess_text(text):
    stop_words = set(stopwords.words('english'))
    lemmatizer = WordNetLemmatizer()
    tokens = word_tokenize(text.lower())
    return " ".join(
        lemmatizer.lemmatize(w) for w in tokens if w.isalnum() and w not in stop_words
    )

def compare_text_similarity(t1, t2):
    embeddings = TEXT_MODEL.encode([t1, t2])
    return round(cosine_similarity([embeddings[0]], [embeddings[1]])[0][0] * 100, 2)

def extract_images_from_pdf(pdf_path, output_folder, prefix):
    doc = fitz.open(pdf_path)
    extracted_imgs, saved_paths = [], []

    for page_no in range(len(doc)):
        images = doc[page_no].get_images(full=True)
        for idx, img in enumerate(images):
            xref = img[0]
            base_image = doc.extract_image(xref)
            img_data = np.frombuffer(base_image["image"], np.uint8)
            img_cv = cv2.imdecode(img_data, cv2.IMREAD_COLOR)

            if img_cv is not None:
                img_cv = cv2.resize(img_cv, (img_cv.shape[1] * 2, img_cv.shape[0] * 2))
                gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
                blurred = cv2.GaussianBlur(gray, (5, 5), 0)
                sharpened = cv2.filter2D(blurred, -1, np.array([[0, -1, 0], [-1, 5, -1], [0, -1, 0]]))
                thresh = cv2.adaptiveThreshold(sharpened, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                               cv2.THRESH_BINARY_INV, 11, 2)
                contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

                if len([c for c in contours if cv2.contourArea(c) > 500]) > 2:
                    filename = f"{prefix}_page{page_no}_img{idx}.png"
                    path = os.path.join(output_folder, filename)
                    cv2.imwrite(path, img_cv)
                    extracted_imgs.append(img_cv)
                    saved_paths.append(path)

    return extracted_imgs, saved_paths

def extract_images_from_docx(docx_path, output_folder, prefix):
    doc = docx.Document(docx_path)
    extracted_imgs, saved_paths = [], []

    for idx, rel in enumerate(doc.part.rels.values()):
        if "image" in rel.reltype:
            img_data = rel.target_part.blob
            img_array = np.frombuffer(img_data, np.uint8)
            img_cv = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

            if img_cv is not None:
                filename = f"{prefix}_img{idx}.png"
                path = os.path.join(output_folder, filename)
                cv2.imwrite(path, img_cv)
                extracted_imgs.append(img_cv)
                saved_paths.append(path)

    return extracted_imgs, saved_paths

def extract_images(path, output_folder, prefix):
    if is_pdf(path):
        return extract_images_from_pdf(path, output_folder, prefix)
    elif is_docx(path):
        return extract_images_from_docx(path, output_folder, prefix)
    return [], []

def extract_text_from_images(images):
    return "\n".join(pytesseract.image_to_string(Image.fromarray(img)) for img in images)

def calculate_grammar_score(text):
    """Return grammar score out of 10 based on simple word-level corrections."""
    blob = TextBlob(text)
    errors = sum(1 for word in blob.words if word != word.correct())
    return 10 if errors == 0 else 9 if errors <= 5 else 8

def calculate_plagiarism_scores(assignment):
    submissions = Submission.objects.filter(assignment=assignment)
    texts = [s.preprocessed_content or "" for s in submissions]
    tfidf = TfidfVectorizer().fit_transform(texts)

    for i, sub in enumerate(submissions):
        similarities = [
            cosine_similarity(tfidf[i], tfidf[j])[0][0]
            for j in range(len(submissions)) if i != j
        ]
        avg_sim = (sum(similarities) / len(similarities)) * 100 if similarities else 0
        sub.plagiarism_score = round(avg_sim, 2)
        sub.save()


# ---------------------- View Functions ----------------------

@login_required
@user_passes_test(lambda u: hasattr(u, 'student_profile'))
def submit_assignment(request, assignment_id):
    assignment = get_object_or_404(Assignment, id=assignment_id)
    student_profile = request.user.student_profile

    if request.method == 'POST':
        form = SubmissionForm(request.POST, request.FILES)
        if form.is_valid():
            submission = form.save(commit=False)
            submission.student = student_profile
            submission.assignment = assignment
            submission.submitted_at = timezone.now()
            submission.is_late = timezone.now() > assignment.due_date

            if submission.is_late and not submission.student_comments:
                messages.error(request, "Since you are submitting after the deadline, you must provide a comment explaining the delay.")
                return render(request, 'submit_assignment.html', {
                    'form': form,
                    'assignment': assignment
                })
        
            submission.save()

            evaluate_submission(submission.id)

            messages.success(request, 'Assignment submitted and evaluated successfully!')
            return redirect('student_dashboard')
    else:
        form = SubmissionForm()

    return render(request, 'submit_assignment.html', {'form': form, 'assignment': assignment})

def evaluate_submission(submission_id):
    sub = Submission.objects.get(id=submission_id)
    assignment = sub.assignment

    # ----- Text Extraction and Saving -----
    student_text = extract_text(sub.file.path)
    teacher_text = extract_text(assignment.model_answer_file.path)

    sub.content = student_text  # ✅ Save extracted raw text to `content`

    preprocessed_student_text = preprocess_text(student_text)
    preprocessed_teacher_text = preprocess_text(teacher_text)
    sub.preprocessed_content = preprocessed_student_text

    # ----- Text Similarity -----
    sub.text_similarity_score = compare_text_similarity(preprocessed_teacher_text, preprocessed_student_text)

    # Save extracted student text into a file inside MEDIA (optional enhancement)
    student_text_path = f"student_texts/submission_{sub.id}.txt"
    full_text_path = os.path.join(settings.MEDIA_ROOT, student_text_path)
    os.makedirs(os.path.dirname(full_text_path), exist_ok=True)
    with open(full_text_path, 'w', encoding='utf-8') as f:
        f.write(student_text)
    sub.text_file = student_text_path  # Assuming you have a `text_file = models.FileField` in your model (optional)

    # ----- Diagram Text Extraction and Saving -----
    teacher_imgs, _ = extract_images(assignment.model_answer_file.path, 'teacher_images', 'teacher')
    student_imgs, _ = extract_images(sub.file.path, 'student_images', 'student')

    teacher_diagram_text = extract_text_from_images(teacher_imgs)
    student_diagram_text = extract_text_from_images(student_imgs)

    sub.image_text = student_diagram_text  # ✅ Save extracted image text into `image_text`

    # Save extracted image text into a file inside MEDIA (optional enhancement)
    student_image_text_path = f"student_image_texts/submission_{sub.id}.txt"
    full_image_text_path = os.path.join(settings.MEDIA_ROOT, student_image_text_path)
    os.makedirs(os.path.dirname(full_image_text_path), exist_ok=True)
    with open(full_image_text_path, 'w', encoding='utf-8') as f:
        f.write(student_diagram_text)
    sub.image_text_file = student_image_text_path  # Assuming you have a `image_text_file = models.FileField` (optional)

    # ----- Image Similarity -----
    image_sim = compare_text_similarity(teacher_diagram_text, student_diagram_text)

    if image_sim > 50:
        sub.image_similarity_score = 100.0  # ✅ Force to 100
    else:
        sub.image_similarity_score = 50.0  # ✅ Force to 50

    # ----- Grammar Score -----
    sub.grammar_score = calculate_grammar_score(preprocessed_student_text)

    # ----- Provisional Marks and Grade (Before Plagiarism) -----
    provisional_score = (
        sub.text_similarity_score * 0.5 +
        sub.image_similarity_score * 0.3 +
        sub.grammar_score * 2
    )
    sub.provisional_marks = round(provisional_score, 2)

    # Set Provisional Grade and Feedback
    for threshold, grade, feedback in [
        (91, "A1", "Awesome start! You're on fire!"),
        (81, "A2", "Great submission! Just a little polish needed."),
        (71, "B1", "Good job! Focus on refining small mistakes."),
        (61, "B2", "Nice effort! More careful work will boost your marks."),
        (51, "C1", "You’re doing okay, but review the concepts again."),
        (41, "C2", "You're halfway there! Keep practicing."),
        (33, "D", "Needs improvement, but you have potential!"),
        (0,  "E", "Please meet your teacher for help. You can do better!"),
    ]:
        if provisional_score >= threshold:
            sub.grade_before_plagiarism = grade  # ✅ Save grade before plagiarism
            sub.feedback_before_plagiarism = feedback  # ✅ Save feedback before plagiarism
            break

    sub.save()



def finalize_submissions_after_deadline(assignment_id):
    assignment = Assignment.objects.get(id=assignment_id)
    submissions = Submission.objects.filter(assignment=assignment)

    calculate_plagiarism_scores(assignment)

    for sub in submissions:
        # Final Grade Calculation after adding plagiarism impact
        score = (
            sub.text_similarity_score * 0.4 +
            sub.image_similarity_score * 0.3 +
            sub.grammar_score * 1 +
            (100 - sub.plagiarism_score) * 0.2
        )
        sub.final_marks = round(score, 2)

        for threshold, grade, feedback in [
            (91, "A1", "Outstanding performance! Keep up the hard work."),
            (81, "A2", "Great job! A little more effort can take you to the top."),
            (71, "B1", "Impressive work! Keep focusing and improving."),
            (61, "B2", "You're on the right track! Practice more to excel."),
            (51, "C1", "Decent effort, but there's room for improvement."),
            (41, "C2", "Fair work, but strive to be better."),
            (33, "D", "Needs more effort. Try to improve next time."),
            (0,  "E", "Poor performance. Please seek help to understand the material."),
        ]:
            if score >= threshold:
                sub.grade_after_plagiarism = grade  # ✅ Save grade after plagiarism
                sub.feedback_after_plagiarism = feedback  # ✅ Save feedback after plagiarism
                break

        sub.save()


def progress_view(request, assignment_id):
    assignment = Assignment.objects.get(id=assignment_id)
    # Assuming you have a method to get the student submissions
    student_submissions = Submission.objects.filter(assignment=assignment)
    
    # Fetch final marks and provisional marks for each student
    student_marks_list = []
    for submission in student_submissions:
        student_marks_list.append({
            'id': submission.student.student.id,
            'name': submission.student.name,
            'provisional_marks': submission.provisional_marks,
            'final_marks': submission.final_marks,
        })
    
    # Pass both provisional and final marks to the new template
    return render(request, 'progress_chart.html', {
        'assignment_title': assignment.title,
        'student_marks_list': student_marks_list,
        'role': 'student',  # Assuming the role is passed in the context
    })


@login_required
def query1to1_view(request, teacher_id, student_id):
    teacher = get_object_or_404(TeacherProfile, teacher__id=teacher_id)
    student = get_object_or_404(StudentProfile, student__id=student_id)

    if request.method == "POST":
        message = request.POST.get("message")
        if message:
            PrivateMessage.objects.create(
                sender=request.user,
                message=message,
                receiver = teacher.teacher if request.user == student.student.student else student.student.student,
                timestamp=now()
            )
        return redirect("query1to1", teacher_id=teacher_id, student_id=student_id)  # Redirect after saving

    messages = PrivateMessage.objects.filter(
        sender__in=[teacher.teacher, student.student],
        receiver__in=[teacher.teacher, student.student]
    ).order_by("-timestamp")  # Get messages between the two users

    return render(request, "query1to1.html", {"teacher": teacher, "student": student, "messages": messages})


@login_required
def queryclassroom_view(request, class_id):
    classroom = get_object_or_404(Classroom, id=class_id)

    if request.method == "POST":
        message = request.POST.get("message")
        print(f"DEBUG: Received message: {message}")  # ✅ Check if message is received

        if message:  # Ensure message is not empty
            query = QueryMessage.objects.create(
                classroom=classroom,
                sender=request.user,
                message=message,
                timestamp=now()
            )
            print(f"DEBUG: Query saved: {query}")  # ✅ Check if query is saved

        return redirect("class_query", class_id=class_id)  # Redirect after saving

    queries = QueryMessage.objects.filter(classroom=classroom).order_by("-timestamp")
    print(f"DEBUG: Queries in DB: {queries}")  # ✅ Check if queries exist

    return render(request, "queryclassroom.html", {"classroom": classroom, "queries": queries})

# Function to notify students of a new assignment
def notify_students_of_new_assignment(assignment_id):
    """Send notifications to students who haven't submitted their assignment yet."""
    assignment = Assignment.objects.get(id=assignment_id)
    students = StudentProfile.objects.all()  # Fetch all students (StudentProfile)

    for student in students:
        # Check if the student has submitted the assignment
        submission = Submission.objects.filter(assignment=assignment, student=student).first()

        # If no submission or not submitted, send a notification
        if not submission or not submission.submitted:
            message = f"You have a new assignment: '{assignment.title}'. Please submit it before the due date."
            create_notification(student, message)  # Create a notification instance


# Notify teacher and student in case of plagiarism
def notify_teacher_and_student(submission, teacher_email, plagiarism_score):
    """Sends an alert email to the teacher and student about plagiarism detection."""
    message = (
        f"Hello,\n\nA submission from {submission.student.student.email} has a plagiarism score of {plagiarism_score:.2f}%.\n"
        f"Assignment: {submission.assignment.title}\n\nPlease review it.\n\nBest regards,\nSubmitTech"
    )
    send_mail(
        "⚠️ Plagiarism Alert: High Similarity Detected!",
        message,
        'noreply@submitech.com',
        [teacher_email, submission.student.student.email]
    )

def send_notifications():
    """Sends reminder emails to students who haven't submitted their assignments."""
    pending_students = CustomUser.objects.filter(role="student", submission__isnull=True).distinct()
    for student in pending_students:
        send_mail(
            '⏳ Assignment Reminder',
            '📌 You have pending assignments. Please submit before the deadline.',
            'admin@submittech.com',
            [student.email],
            fail_silently=True)

# Function to create a new notification
def create_notification(student, message):
    """Create a notification entry in the database."""
    Notification.objects.create(student=student, message=message)

# Function to get unread notifications for a student
def get_notifications(student):
    """Get all unread notifications for a student."""
    return Notification.objects.filter(student=student, is_read=False)

# Function to mark a notification as read
def mark_notification_as_read(notification_id):
    """Mark a notification as read."""
    notification = Notification.objects.get(id=notification_id)
    notification.is_read = True
    notification.save()